# Header Files
from docx import Document
import pandas as pd
from googletrans import Translator
import httpx

timeout = httpx.Timeout(10)
translator = Translator(timeout=timeout)

# File imports
converted_document_path = './Converted -  Question Paper.docx'
document_path = './Question Paper.docx'
converted = Document(converted_document_path)
document = Document(document_path)


# Import Questions Data
tables = []
for table in document.tables:
    df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text:
                df[i][j] = cell.text
    tables.append(pd.DataFrame(df))


# Translated Questions
dict = {'Question': [], 'Answer': []}
for i in range(len(tables[0])):
    dict['Question'].append("{0}) ".format(i + 1) + tables[0][0][i])
    dict['Answer'].append(tables[0][1][i])
    dict['Question'].append(translator.translate("{0}) ".format(i + 1) + tables[0][0][i], src='en', dest='te').text)
    dict['Answer'].append(tables[0][1][i])
df = pd.DataFrame(dict)


# Writing the Converted Questions in document
t = converted.add_table(df.shape[0] + 1, df.shape[1])
for j in range(df.shape[-1]):
    t.cell(0, j).text = df.columns[j]
for i in range(df.shape[0]):
    for j in range(df.shape[-1]):
        t.cell(i + 1, j).text = str(df.values[i, j])


# Saving the Document
converted.save(converted_document_path)
